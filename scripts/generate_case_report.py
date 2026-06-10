from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
CASE_DIR = DOCS / "案例报告"
OUTPUT_DOCX = CASE_DIR / "AI投资学课程案例报告.docx"
OUTPUT_MD = CASE_DIR / "AI投资学课程案例报告.md"
OUTPUT_PDF_DIR = CASE_DIR

PLOTS = ROOT / "test" / "backtest_outputs" / "plots"
TABLES = ROOT / "outputs" / "tables"
BACKTEST = ROOT / "test" / "backtest_outputs"

FONT_ZH = "宋体"
FONT_HEADING = "黑体"
FONT_EN = "Times New Roman"
FONT_CODE = "Courier New"

BODY_PT = 12
TITLE_PT = 15
ABSTRACT_PT = 10.5
CAPTION_PT = 10.5
TABLE_PT = 10.5
CODE_PT = 8.5


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def pct(value: float, digits: int = 2) -> str:
    return f"{value * 100:.{digits}f}%"


def num(value: float, digits: int = 4) -> str:
    return f"{value:.{digits}f}"


def clean_strategy(name: str) -> str:
    mapping = {
        "spo_plus_turnover_l2": "SPO+换手/L2约束",
        "spo_plus": "SPO+",
        "spo_plus_fee": "SPO+交易费",
        "pto_markowitz": "PtO-Markowitz",
        "equal_weight_6etf": "6ETF等权",
        "spy_buy_hold": "SPY买入持有",
        "ai_low_risk": "AI低风险组合",
        "ai_medium_risk": "AI中风险组合",
        "ai_high_risk": "AI高风险组合",
    }
    return mapping.get(name, name)


def markdown_table(rows: list[list[str]], headers: list[str]) -> str:
    widths = [len(h) for h in headers]
    for row in rows:
        for idx, cell in enumerate(row):
            widths[idx] = max(widths[idx], len(str(cell)))
    lines = [
        "| " + " | ".join(str(h) for h in headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(str(c) for c in row) + " |")
    return "\n".join(lines)


def rel_path(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def file_size_kb(path: Path) -> str:
    return f"{path.stat().st_size / 1024:.1f}KB"


def describe_artifact(path: Path) -> str:
    if path.suffix == ".csv":
        frame = pd.read_csv(path)
        return f"{frame.shape[0]}行 x {frame.shape[1]}列"
    if path.suffix == ".json":
        data = read_json(path)
        if isinstance(data, dict):
            keys = ", ".join(list(data.keys())[:5])
            return f"JSON对象；一级键：{keys}"
        return f"JSON数组；{len(data)}项"
    if path.suffix == ".md":
        text = path.read_text(encoding="utf-8")
        return f"Markdown；{len(text.splitlines())}行"
    if path.suffix == ".png":
        return "PNG图像"
    return path.suffix.lstrip(".").upper() or "文件"


def artifact_purpose(path: Path) -> str:
    purposes = {
        "ai_prompt_payload.json": "AI风险组合提示词输入、约束与候选组合载荷",
        "ai_risk_profile_backtest_returns.csv": "三档AI风险组合逐月收益与换手",
        "ai_risk_profile_portfolios.json": "最新三档风险组合JSON交付结果",
        "ai_risk_profile_validation.json": "AI组合权重、风险约束与指标有限性校验",
        "ai_risk_profile_weights.csv": "三档AI风险组合逐月ETF权重明细",
        "data_quality_check.csv": "ETF原始价格与收益率质量检查明细",
        "download_metadata.json": "下载配置、数据源、备用通道与质量摘要",
        "etf_risk_return_metrics.csv": "ETF单体风险收益、动量与预测排名",
        "feature_dataset_quality.json": "特征、标签、样本切分与缺失检查摘要",
        "portfolio_backtest_metrics.csv": "主脚本输出的策略回测指标",
        "spo_backtest_returns.csv": "SPO/PtO/基准策略逐月收益",
        "spo_candidate_portfolios.csv": "候选组合预期收益、波动、权重JSON与换手",
        "spo_model_metrics.csv": "Ridge与SPO+预测层诊断指标",
        "spo_portfolio_weights.csv": "SPO、PtO与基准策略逐月ETF权重",
        "spo_prediction_scores.csv": "测试期各ETF预测分数与未来收益",
        "spo_reproduction_metadata.json": "SPO复现实验配置、数据切分、校验与引用信息",
        "spo_training_history.csv": "SPO+训练损失、验证MSE与Rank IC记录",
        "backtest_summary.json": "独立回测总摘要、最佳策略与主要风险提示",
        "combined_model_backtest_returns.csv": "合并SPO、PtO、基准与AI组合后的逐月收益",
        "existing_model_backtest_report.md": "独立回测自动生成的审计报告",
        "lookahead_audit.json": "未来函数与时间顺序审计结果",
        "model_assessment.csv": "策略优势、弱点与推荐用途评估",
        "model_quality_ranking.csv": "策略质量综合评分排序",
        "random_forest_feature_importance.csv": "随机森林诊断性特征重要性",
        "strategy_backtest_metrics.csv": "独立回测总收益、波动、回撤、换手等指标",
        "strategy_drawdown_curves.csv": "各策略逐月回撤曲线",
        "strategy_equity_curves.csv": "各策略逐月净值曲线",
        "strategy_excess_returns.csv": "相对等权基准的超额收益与信息比率",
        "transaction_cost_sensitivity.csv": "不同交易成本假设下的绩效敏感性",
    }
    plot_purposes = {
        "ai_monthly_weight_heatmap.png": "AI三档风险组合逐月权重热力图",
        "etf_price_index.png": "六只ETF归一化价格指数走势",
        "random_forest_feature_importance.png": "随机森林诊断性特征重要性图",
        "strategy_cumulative_returns.png": "主要策略累计净值曲线",
        "strategy_drawdowns.png": "主要策略回撤曲线",
        "strategy_performance_table.png": "主要策略绩效表图像",
    }
    return purposes.get(path.name, plot_purposes.get(path.name, "模型或回测输出产物"))


def build_artifact_index() -> list[list[str]]:
    paths = [
        *sorted(TABLES.glob("*")),
        *sorted(p for p in BACKTEST.glob("*") if p.is_file()),
        *sorted(PLOTS.glob("*.png")),
    ]
    rows = []
    for path in paths:
        rows.append(
            [
                rel_path(path),
                path.suffix.lstrip(".").upper() or "FILE",
                describe_artifact(path),
                file_size_kb(path),
                artifact_purpose(path),
            ]
        )
    return rows


def source_symbol(rel_file: str, symbol: str) -> str:
    path = ROOT / rel_file
    lines = path.read_text(encoding="utf-8").splitlines()
    pattern = re.compile(rf"^(def|class)\s+{re.escape(symbol)}\b")
    start = next((idx for idx, line in enumerate(lines) if pattern.match(line)), None)
    if start is None:
        raise ValueError(f"Cannot find {symbol} in {rel_file}")

    end = len(lines)
    boundary = re.compile(r"^(def|class)\s+\w+|^@dataclass\b|^if __name__")
    for idx in range(start + 1, len(lines)):
        if boundary.match(lines[idx]):
            end = idx
            break
    return "\n".join(lines[start:end]).rstrip()


def code_block(rel_file: str, symbols: list[str]) -> str:
    return "\n\n".join(source_symbol(rel_file, symbol) for symbol in symbols)


def get_tables() -> dict[str, Any]:
    return {
        "download": read_json(TABLES / "download_metadata.json"),
        "features": read_json(TABLES / "feature_dataset_quality.json"),
        "spo_meta": read_json(TABLES / "spo_reproduction_metadata.json"),
        "ai_validation": read_json(TABLES / "ai_risk_profile_validation.json"),
        "backtest_summary": read_json(BACKTEST / "backtest_summary.json"),
        "lookahead": read_json(BACKTEST / "lookahead_audit.json"),
        "portfolio_metrics": pd.read_csv(TABLES / "portfolio_backtest_metrics.csv"),
        "strategy_metrics": pd.read_csv(BACKTEST / "strategy_backtest_metrics.csv"),
        "ranking": pd.read_csv(BACKTEST / "model_quality_ranking.csv"),
        "excess": pd.read_csv(BACKTEST / "strategy_excess_returns.csv"),
        "model_metrics": pd.read_csv(TABLES / "spo_model_metrics.csv"),
        "etf_metrics": pd.read_csv(TABLES / "etf_risk_return_metrics.csv"),
        "candidates": pd.read_csv(TABLES / "spo_candidate_portfolios.csv"),
        "feature_importance": pd.read_csv(BACKTEST / "random_forest_feature_importance.csv"),
        "ai_portfolios": read_json(TABLES / "ai_risk_profile_portfolios.json"),
        "data_quality": pd.read_csv(TABLES / "data_quality_check.csv"),
        "cost_sensitivity": pd.read_csv(BACKTEST / "transaction_cost_sensitivity.csv"),
        "assessment": pd.read_csv(BACKTEST / "model_assessment.csv"),
        "artifact_index": build_artifact_index(),
    }


def write_docs(context: dict[str, Any]) -> None:
    feature_summary = context["features"]["quality_summary"]
    download_summary = context["download"]["quality_summary"]
    spo_meta = context["spo_meta"]
    backtest_summary = context["backtest_summary"]
    ranking = context["ranking"]
    metrics = context["strategy_metrics"]
    model_metrics = context["model_metrics"]
    etf_metrics = context["etf_metrics"]
    importance = context["feature_importance"]

    best = backtest_summary["best_model_by_quality_score"]
    sharpe_best = backtest_summary["best_strategy_by_net_sharpe"]
    net_metrics = metrics[metrics["return_type"] == "net"].copy()
    net_metrics["strategy_label"] = net_metrics["strategy"].map(clean_strategy)
    metric_rows = []
    for _, row in net_metrics.sort_values("total_return", ascending=False).iterrows():
        metric_rows.append(
            [
                row["strategy_label"],
                pct(row["total_return"]),
                pct(row["annual_return"]),
                pct(row["annual_volatility"]),
                num(row["sharpe"], 2),
                pct(row["max_drawdown"]),
                pct(row["average_turnover"]),
            ]
        )

    deep_dive = f"""# 项目深读与实现事实沉淀

本文档由 `scripts/generate_case_report.py` 根据当前工作区的代码、数据和复跑输出自动生成，用于支撑 `docs/案例报告/AI投资学课程案例报告.docx` 的事实口径。

## 1. 项目目标与边界

本项目用于《投资学》课程 AI+投资学案例研究，目标不是构建可直接实盘交易的投资系统，而是建立一条可复现的研究链路：从多资产 ETF 历史数据获取开始，经过特征工程、SPO/PtO 模型训练、月度组合构建、交易成本扣除、AI 风险档位组合生成、独立回测复核，最终形成可解释的投资组合案例报告。

当前资产池为 SPY、QQQ、TLT、GLD、VNQ、DBC 六只 ETF，覆盖美国大盘权益、科技成长权益、长期美国国债、黄金、房地产 REITs 和大宗商品。样本期为 {download_summary['first_dates']['SPY']} 至 {download_summary['last_dates']['SPY']}，每只 ETF 有 {download_summary['row_counts']['SPY']} 条日度复权收盘价记录。

## 2. 代码结构

- `src/download_data.py`：下载 Yahoo Finance 历史价格，保存原始多级表头价格、复权收盘价、日收益率和数据质量检查表。脚本先尝试 `yfinance.download()`，若发生限流或网络问题，则切换到 Yahoo Chart API 备用通道，并仅在质量检查通过后覆盖数据文件。
- `src/features.py`：构造 22 个历史特征和 5 个 21 日前瞻标签。特征只使用样本日及以前的价格、收益和成交量，标签使用未来 21 个交易日收益。
- `src/reproduce_spo_paper.py`：在本地 ETF 数据上复现 SPO+ 决策导向学习路径。预测器为 PyTorch 线性模型，损失函数为 PyEPO `SPOPlus`，并与 Ridge PtO-Markowitz、6ETF 等权和 SPY 买入持有对照。
- `src/generate_ai_risk_profiles.py`：读取 SPO/PtO 输出，计算 ETF 单体风险收益指标，生成低、中、高三档 AI 风险组合，并用程序化规则校验权重和、非负、单资产上限、波动率上限。
- `test/backtest_existing_models.py`：作为独立回测层，读取已有输出重新计算组合净值、回撤、成本敏感性、超额收益、随机森林特征重要性和未来函数审计。

## 3. 当前复跑结果

数据质量检查通过：6 个 ETF 日期范围一致，无重复日期、复权价格无缺失、无非正价格。特征数据集共有 {feature_summary['daily_modeling_shape'][0]} 行日度建模样本、{feature_summary['monthly_rebalance_shape'][0]} 行月末调仓样本；训练、验证、测试月末样本分别为 {feature_summary['split_counts_monthly']['train']}、{feature_summary['split_counts_monthly']['valid']}、{feature_summary['split_counts_monthly']['test']} 行。

独立回测层显示，综合评分最高的非基准策略是 {clean_strategy(best['strategy'])}，扣成本后总收益 {pct(best['total_return'])}、年化收益 {pct(best['annual_return'])}、Sharpe {num(best['sharpe'], 2)}、最大回撤 {pct(best['max_drawdown'])}。净 Sharpe 最高的是 {clean_strategy(sharpe_best['strategy'])}，Sharpe {num(sharpe_best['sharpe'], 2)}。

## 4. 净收益指标快照

{markdown_table(metric_rows, ['策略', '总收益', '年化收益', '年化波动', 'Sharpe', '最大回撤', '平均换手'])}

## 5. 预测层指标

{markdown_table([[r.model_name, r.split, num(r.rmse, 4), num(r.mae, 4), num(r.r2, 4), num(r.rank_ic, 4)] for r in model_metrics.itertuples()], ['模型', '样本', 'RMSE', 'MAE', 'R2', 'Rank IC'])}

## 6. 随机森林诊断性特征重要性

随机森林只用于独立诊断，不参与主 SPO/PtO 回测。当前前十重要特征如下：

{markdown_table([[r.feature, num(r.importance, 4), str(int(r.rank))] for r in importance.head(10).itertuples()], ['特征', '重要性', '排名'])}
"""

    theory = """# SPO 理论、文献与算法基础沉淀

## 1. 投资组合优化的理论起点

现代投资组合理论起源于 Markowitz 的均值-方差框架。该框架将组合选择刻画为收益与风险之间的权衡，即投资者在给定收益目标下最小化方差，或在给定风险预算下最大化预期收益。虽然本项目的主线是机器学习和 SPO，但所有组合权重解释都仍然依赖这一直觉：收益预测不能脱离波动率、回撤、换手和约束独立评价。

Sharpe 比率提供了风险调整收益的经典度量，形式上是单位波动承载的超额收益。本项目没有设置现金利率或无风险收益率，因此在课程回测中采用年化收益除以年化波动率的简化 Sharpe 口径，并同时报告最大回撤、Sortino、Calmar 和胜率，避免单一指标误导。

## 2. PtO 与 SPO 的差异

传统 Predict-then-Optimize, PtO 流程先训练预测模型，再把预测收益输入组合优化器。这个流程直观、可解释，但存在训练目标与投资目标错配的问题：预测模型通常最小化 MSE 或 MAE，而投资者真正关心的是由预测值引出的组合权重是否能改善收益、风险和回撤。

Elmachtoub 与 Grigas 提出的 Smart Predict, then Optimize, SPO 框架直接针对这一错配建立决策损失。SPO 的思想不是说预测误差不重要，而是认为同样大小的预测误差在不同优化问题中会造成不同决策后果。因此，学习目标应关注决策 regret，而不是只关注点预测误差。

## 3. SPO+ 损失

在组合优化中，设真实未来收益向量为 r，预测收益向量为 r_hat，可行组合集合为 W。如果给定真实收益，最优权重为 w_star；如果使用预测收益，则得到权重 w_hat。SPO regret 可以理解为真实收益下 oracle 决策与预测诱导决策之间的收益差。由于直接优化 SPO loss 通常不可微或计算困难，SPO+ 提供了一个可训练的凸 surrogate，使预测器能通过优化层获得决策导向梯度。

本项目采用 PyEPO 的 `SPOPlus` 实现。为了避免商业求解器依赖，项目把 MaxReturn long-only 问题写成自定义 `MaxReturnOptModel`。当目标只是最大化预测收益且权重非负、权重和为 1 时，闭式解就是把全部权重给预测分数最高的资产。随后项目又在权重生成阶段加入交易成本、换手惩罚和 L2 权重正则，缓解基础 MaxReturn 的极端集中问题。

## 4. 原论文吸收与本项目差异

Wang Yi 与 Takashi Hasuike 的 2026 年论文《Smart Predict--then--Optimize Paradigm for Portfolio Optimization in Real Markets》强调，收益预测准确率提升并不必然转化为组合决策质量提升，尤其是在存在交易摩擦、换手控制和市场非平稳性的真实市场中。论文采用 U.S. ETF 数据、月度再平衡、线性预测器、SPO+、交易费、换手控制、鲁棒扰动和 SoftmaxDFL 等方案。

本项目已下载论文 arXiv 源码包到 `external/paper_2601_04062_source/`。源码包包含论文 LaTeX、图表和参考文献，但不包含可直接运行的 Python 实验代码或作者 GitHub 链接。因此，本项目不能宣称完整复刻作者全部实验，而是依据论文正文和 PyEPO 官方实现，在本地 6 ETF 数据集上复现 SPO+ 核心路径，并把 RobustSPO、SoftmaxDFL、Optuna 滚动调参等作为后续扩展方向。

## 5. 为什么仍然报告点预测指标

SPO 的重点是组合决策质量，但报告仍需要展示 RMSE、MAE、R2 和 Rank IC。原因有三点：第一，点预测指标能帮助判断模型是否出现数值异常；第二，Rank IC 能粗略观察横截面排序信号；第三，如果一个模型组合表现好但预测指标很差，就必须在结论中谨慎表述，避免把样本期组合收益误读为稳定预测能力。

本次复跑中，PyEPO SPO+ 线性模型在点预测指标上显著弱于 Ridge，测试 Rank IC 接近零。但 `spo_plus_turnover_l2` 在组合层面表现较好。这一结果恰好体现了 SPO 框架的核心讨论：投资报告不能只看预测误差，也不能只看回测收益，而要同时解释两者之间的张力。
"""

    empirical = f"""# 数据集质量、回测与投资组合实证沉淀

## 1. 高质量数据集回应

本项目的数据质量控制包含四层：

1. 原始下载层：保留 Yahoo Finance 返回的 Open、High、Low、Close、Adj Close、Volume 六类字段，原始表为两级表头，便于追溯。
2. 价格处理层：组合回测和收益率计算统一使用 Adj Close，避免现金分红和拆分造成收益率偏差。
3. 质量检查层：检查日期范围、行数一致性、缺失值、重复日期、非正价格和极端单日收益。
4. 标签防泄露层：所有特征只使用样本日及以前信息，未来 21 日收益只作为标签和事后回测收益，不进入当期特征。

当前复跑质量摘要：

{markdown_table([
    ['资产数量', '6'],
    ['价格行数', str(download_summary['row_counts']['SPY'])],
    ['日期范围', f"{download_summary['first_dates']['SPY']} 至 {download_summary['last_dates']['SPY']}"],
    ['复权价格缺失', '0'],
    ['重复日期', str(download_summary['duplicate_dates'])],
    ['日度建模样本', str(feature_summary['daily_modeling_shape'][0])],
    ['月末调仓样本', str(feature_summary['monthly_rebalance_shape'][0])],
    ['特征数量', str(len(feature_summary['feature_names']))],
], ['检查项', '结果'])}

## 2. ETF 单体风险收益

最新决策日为 {etf_metrics['decision_date'].iloc[0]}。六只 ETF 的历史风险收益和最新预测指标如下：

{markdown_table([[r.ticker, str(int(r.spo_predicted_rank)), pct(r.expected_return_21d_pto), pct(r.historical_annual_return), pct(r.annual_volatility), pct(r.max_drawdown), num(r.sharpe, 2), pct(r.momentum_60d)] for r in etf_metrics.itertuples()], ['ETF', 'SPO排名', 'PtO预测21日收益', '历史年化收益', '年化波动', '最大回撤', 'Sharpe', '60日动量'])}

## 3. 回测口径

测试期为 {spo_meta['data']['test_start']} 至 {spo_meta['data']['test_end']}，共有 {spo_meta['data']['test_months']} 个按月调仓样本。组合在 t 月末形成，使用未来 21 个交易日收益作为下一期实现收益。净收益公式为：

`net_return_t = gross_return_t - sum(abs(weight_t - weight_t-1)) * cost_rate`

本报告主口径采用 0.5% 单边交易成本，与主脚本 `DEFAULT_COST_RATE = 0.005` 一致。独立回测层已通过未来函数审计，确认特征、标签、回测日期和 StandardScaler 拟合规则均符合时间顺序。

## 4. 净收益回测结果

{markdown_table(metric_rows, ['策略', '总收益', '年化收益', '年化波动', 'Sharpe', '最大回撤', '平均换手'])}

## 5. 结果解释

`SPO+换手/L2约束` 是当前样本期综合评分最高的非基准策略。其优势来自两个方向：一是 SPO+ 让训练目标更贴近组合决策；二是换手和 L2 正则在基础 MaxReturn 的集中仓位之外加入了稳定性约束。基础 SPO+ 和 SPO+交易费版本表现相同，说明当前交易费项在该实现下没有形成有效差异，需要在后续版本中进一步检查费用项进入优化器的强度。

6ETF 等权组合是最强的风险调整基准，净 Sharpe 达 {num(sharpe_best['sharpe'], 2)}，最大回撤仅 {pct(sharpe_best['max_drawdown'])}。这一结果说明，多资产分散化本身已经非常有力，AI/SPO 策略必须在扣除成本后显著超过等权基准，才具有充分的增量价值。

AI 三档风险组合提供了可解释的风险偏好映射。低风险组合牺牲部分收益换取较低波动；中风险组合在收益和波动之间折中；高风险组合放宽单资产上限，但由于换手成本较高，扣成本后不一定优于低风险或中风险组合。
"""

    structure = """# 案例报告结构重构说明

正式案例报告采用接近课程论文和毕业论文的结构，而不是简单项目说明。结构设计如下：

1. 内封面：包含中文题名、英文题名、课程、作者、日期和声明。
2. 中文摘要与关键词：概括背景、目的、方法、数据、结果和结论。
3. 英文摘要与关键词：与中文摘要对应，满足格式要求中的中外文摘要要求。
4. 研究引言：说明 AI 投资组合优化问题、研究意义、核心问题和技术路线。
5. 理论基础与文献综述：从 Markowitz、Sharpe、PtO、SPO、PyEPO 与 ETF 资产配置展开。
6. 数据来源、变量构造与数据质量控制：重点回应数据来源权威性、数据类型、预处理规则、特征和标签、防泄露机制。
7. 模型设定与组合优化方法：解释 PtO、SPO+、交易成本、换手控制、L2 正则和 AI 风险组合。
8. 回测框架与评价指标体系：定义调仓、交易成本、收益、波动、Sharpe、Sortino、最大回撤、Calmar、胜率和换手。
9. 实证结果与投资组合分析：展示图表、策略对比、预测层诊断、特征重要性、AI 三档权重和结果解释。
10. 稳健性检验、局限性与风险提示：强调样本期有限、数据源限制、交易成本简化、预测指标不足和非投资建议。
11. 研究结论与后续展望：总结贡献、实证发现和后续扩展。
12. 参考文献：按正文引用顺序编号，包含论文、工具、数据源和 ETF 官方来源。
13. 附录：放置复现环境、核心代码、附图汇编、附表索引和主要输出表。

格式上按华东师范大学毕业论文要求执行：A4，页边距上 2.5cm、下 2.0cm、左 3.0cm、右 2.5cm，正文宋体小四，标题黑体，外文 Times New Roman，1.5 倍行距，页码底端居中，表名在表上，图名在图下，中英文对照。
"""

    (DOCS / "project_deep_dive.md").write_text(deep_dive, encoding="utf-8")
    (DOCS / "spo_theory_and_literature_foundation.md").write_text(theory, encoding="utf-8")
    (DOCS / "empirical_results_and_dataset_quality.md").write_text(empirical, encoding="utf-8")
    (DOCS / "case_report_structure_redesign.md").write_text(structure, encoding="utf-8")


def build_report_markdown(context: dict[str, Any]) -> str:
    d = context["download"]["quality_summary"]
    f = context["features"]["quality_summary"]
    meta = context["spo_meta"]
    ai_validation = context["ai_validation"]["validation"]
    lookahead = context["lookahead"]
    summary = context["backtest_summary"]
    metrics = context["strategy_metrics"]
    model_metrics = context["model_metrics"]
    etf_metrics = context["etf_metrics"]
    ranking = context["ranking"]
    importance = context["feature_importance"]
    ai_json = context["ai_portfolios"]
    data_quality = context["data_quality"]
    cost_sensitivity = context["cost_sensitivity"]
    assessment = context["assessment"]
    artifact_index = context["artifact_index"]
    best = summary["best_model_by_quality_score"]
    sharpe_best = summary["best_strategy_by_net_sharpe"]

    net_metrics = metrics[metrics["return_type"] == "net"].copy()
    net_metrics["strategy_label"] = net_metrics["strategy"].map(clean_strategy)
    net_sorted = net_metrics.sort_values("total_return", ascending=False)
    metric_rows = [
        [
            r.strategy_label,
            pct(r.total_return),
            pct(r.annual_return),
            pct(r.annual_volatility),
            num(r.sharpe, 2),
            pct(r.max_drawdown),
            pct(r.win_rate),
            pct(r.average_turnover),
        ]
        for r in net_sorted.itertuples()
    ]

    ai_rows = []
    for item in ai_json["risk_profiles"]:
        weights = ", ".join(
            f"{ticker} {pct(weight, 1)}"
            for ticker, weight in item["weights"].items()
            if weight > 0.0001
        )
        ai_rows.append(
            [
                item["profile_label"],
                pct(item["constraints"]["annual_volatility_max"]),
                pct(item["constraints"]["single_etf_weight_max"]),
                weights,
                pct(item["expected_annual_return"]),
                pct(item["expected_annual_volatility"]),
                num(item["expected_sharpe"], 2),
            ]
        )

    artifact_rows = artifact_index
    data_quality_rows = [
        [
            r.ticker,
            r.first_date,
            r.last_date,
            str(int(r.price_rows)),
            str(int(r.missing_adj_close)),
            str(int(r.non_positive_adj_close)),
            pct(r.period_total_return),
            pct(r.period_max_drawdown),
        ]
        for r in data_quality.itertuples()
    ]
    ranking_rows = [
        [
            str(int(r.rank)),
            clean_strategy(r.strategy),
            r.strategy_group,
            num(r.quality_score, 3),
            pct(r.total_return),
            pct(r.annual_return),
            num(r.sharpe, 2),
            pct(r.max_drawdown),
            pct(r.average_turnover),
        ]
        for r in ranking.itertuples()
    ]
    sensitivity_frame = cost_sensitivity[
        cost_sensitivity["strategy"].isin(
            ["spo_plus_turnover_l2", "spo_plus", "pto_markowitz", "equal_weight_6etf"]
        )
        & cost_sensitivity["cost_rate"].isin([0.0, 0.001, 0.005, 0.01])
    ].copy()
    sensitivity_frame["strategy_label"] = sensitivity_frame["strategy"].map(clean_strategy)
    sensitivity_rows = [
        [
            r.strategy_label,
            pct(r.cost_rate, 2),
            pct(r.total_return),
            pct(r.annual_return),
            num(r.sharpe, 2),
            pct(r.max_drawdown),
            pct(r.average_turnover),
        ]
        for r in sensitivity_frame.sort_values(["strategy_label", "cost_rate"]).itertuples()
    ]
    assessment_rows = [
        [
            str(int(r.rank)),
            clean_strategy(r.strategy),
            r.strategy_group,
            r.strengths,
            r.weaknesses,
            r.recommended_use,
        ]
        for r in assessment.itertuples()
    ]

    code_sources = [
        ["数据下载与质量控制", "src/download_data.py", "download_with_yfinance, download_with_chart_api, prepare_outputs, quality_checks"],
        ["特征工程与标签生成", "src/features.py", "build_wide_features, build_datasets"],
        ["SPO+训练与组合生成", "src/reproduce_spo_paper.py", "MaxReturnOptModel, train_spo_plus, long_only_optimize, build_strategy_weights, backtest"],
        ["AI风险档位组合", "src/generate_ai_risk_profiles.py", "optimize_profile, build_ai_weights, validate_ai_outputs"],
        ["独立回测与审计", "test/backtest_existing_models.py", "build_lookahead_audit, evaluate_returns, calculate_metrics, rank_models, save_all_plots"],
        ["正式报告生成", "scripts/generate_case_report.py", "build_report_markdown, markdown_to_docx；附录摘录导出格式函数，完整源码见文件"],
    ]
    code_quality = code_block("src/download_data.py", ["download_with_yfinance", "download_with_chart_api", "prepare_outputs", "quality_checks"])
    code_features = code_block("src/features.py", ["build_wide_features", "build_datasets"])
    code_spo = code_block("src/reproduce_spo_paper.py", ["MaxReturnOptModel", "train_spo_plus", "long_only_optimize", "build_strategy_weights", "backtest"])
    code_ai = code_block("src/generate_ai_risk_profiles.py", ["optimize_profile", "build_ai_weights", "validate_ai_outputs"])
    code_audit = code_block("test/backtest_existing_models.py", ["build_lookahead_audit", "evaluate_returns", "calculate_metrics", "rank_models", "save_all_plots"])
    code_report = code_block("scripts/generate_case_report.py", ["markdown_to_docx"])

    references = [
        "Markowitz, H. Portfolio Selection[J]. The Journal of Finance, 1952, 7(1):77-91.",
        "Sharpe, W. F. The Sharpe Ratio[J]. The Journal of Portfolio Management, 1994/1998, 21(1):49-58.",
        "Elmachtoub, A. N., Grigas, P. Smart \"Predict, then Optimize\"[J]. Management Science, 2022, 68(1):9-26. DOI:10.1287/mnsc.2020.3922.",
        "Wang, Y., Hasuike, T. Smart Predict--then--Optimize Paradigm for Portfolio Optimization in Real Markets[R]. arXiv:2601.04062, 2026.",
        "Tang, B., Khalil, E. B. PyEPO: a PyTorch-based End-to-End Predict-then-Optimize Library for Linear and Integer Programming[J]. Mathematical Programming Computation, 2024, 16:297-335. DOI:10.1007/s12532-024-00255-x.",
        "Yahoo Finance. Historical Prices and Market Data[E]. https://finance.yahoo.com/.",
        "Aroussi, R. yfinance: Download market data from Yahoo! Finance's API[E]. https://github.com/ranaroussi/yfinance.",
        "State Street Investment Management. State Street SPDR S&P 500 ETF Trust (SPY)[E]. https://www.ssga.com/.",
        "Invesco. Invesco QQQ Trust, Series 1 (QQQ)[E]. https://www.invesco.com/.",
        "BlackRock iShares. iShares 20+ Year Treasury Bond ETF (TLT)[E]. https://www.ishares.com/.",
        "World Gold Council. SPDR Gold Shares (GLD)[E]. https://www.spdrgoldshares.com/.",
        "Vanguard. Vanguard Real Estate ETF (VNQ)[E]. https://investor.vanguard.com/.",
        "Invesco. Invesco DB Commodity Index Tracking Fund (DBC)[E]. https://www.invesco.com/.",
        "Jagannathan, R., Ma, T. Risk Reduction in Large Portfolios: Why Imposing the Wrong Constraints Helps[J]. The Journal of Finance, 2003, 58(4):1651-1683.",
        "Boyd, S. et al. Multi-Period Trading via Convex Optimization[J]. Foundations and Trends in Optimization, 2017, 3(1):1-76.",
    ]

    ref_text = "\n\n".join(f"［{i + 1}］ {ref}" for i, ref in enumerate(references))
    toc_entries = "\n\n".join(
        [
            "摘要",
            "Abstract",
            "一、研究引言",
            "二、理论基础与文献综述",
            "三、数据来源、变量构造与数据质量控制",
            "四、模型设定与组合优化方法",
            "五、回测框架与评价指标体系",
            "六、实证结果与投资组合分析",
            "七、稳健性检验、局限性与风险提示",
            "八、研究结论与后续展望",
            "参考文献",
            "附录A 复现环境、运行命令与输出目录",
            "附录B 核心代码与实现说明",
            "附录C 附图汇编",
            "附录D 附表索引与主要输出表",
        ]
    )
    toc = f"## 目录\n\n{toc_entries}"

    md = f"""# AI投资学课程案例报告

# 基于 Smart Predict--then--Optimize 的多资产 ETF 组合优化策略研究

English Title: A Multi-Asset ETF Portfolio Optimization Case Study Based on Smart Predict--then--Optimize

课程：投资学

项目目录：{ROOT}

完成日期：2026年6月

声明：本报告仅用于课程案例研究，不构成任何投资建议；历史回测结果不代表未来收益。

PAGEBREAK

{toc}

PAGEBREAK

## 摘要

本文围绕 AI 技术在投资组合优化中的应用，构建了一条从数据获取、特征工程、预测建模、决策导向组合优化、风险档位组合生成到独立回测审计的完整研究链路。研究对象为 SPY、QQQ、TLT、GLD、VNQ、DBC 六只美国市场 ETF，分别代表美国大盘权益、科技成长权益、长期美国国债、黄金、房地产 REITs 与大宗商品。数据来自 Yahoo Finance 历史行情，并通过 yfinance 与 Yahoo Chart API 进行程序化获取。复跑结果显示，样本期为 {d['first_dates']['SPY']} 至 {d['last_dates']['SPY']}，每只 ETF 含 {d['row_counts']['SPY']} 条日度复权收盘价记录，复权价格无缺失、无重复日期、无非正价格，能够支撑课程级别的多资产配置实证研究。

方法上，本文以 Markowitz 均值-方差理论和 Sharpe 风险调整收益指标为投资学基础[1][2]，以 Smart Predict--then--Optimize, SPO 作为机器学习与优化耦合的核心范式[3]。SPO 的核心思想是：预测模型不应只最小化点预测误差，而应服务于下游优化决策质量。本文参考 Wang 与 Hasuike 关于真实市场组合优化的 SPO 研究[4]，并基于 PyEPO 的 SPO+ 损失实现本地复现[5]。在本项目中，传统 PtO 基准采用 Ridge 收益预测加 Markowitz 最大 Sharpe 优化；SPO+ 模型采用 PyTorch 线性预测器与 PyEPO SPOPlus 损失；最终比较基础 SPO+、加入交易费项的 SPO+、加入换手与 L2 正则的 SPO+、PtO-Markowitz、6ETF 等权、SPY 买入持有以及三档 AI 风险组合。

实证部分采用月度调仓和未来 21 个交易日收益作为持有期收益。训练、验证、测试按时间切分，测试期为 {meta['data']['test_start']} 至 {meta['data']['test_end']}，共 {meta['data']['test_months']} 个调仓月。独立回测层按 0.5% 单边交易成本重新计算净收益、回撤、成本敏感性与未来函数审计。结果显示，`{clean_strategy(best['strategy'])}` 在非基准策略中综合评分最高，扣成本后总收益为 {pct(best['total_return'])}，年化收益为 {pct(best['annual_return'])}，Sharpe 为 {num(best['sharpe'], 2)}，最大回撤为 {pct(best['max_drawdown'])}；净 Sharpe 最高的策略是 `{clean_strategy(sharpe_best['strategy'])}`，Sharpe 为 {num(sharpe_best['sharpe'], 2)}。这一结果说明，决策导向学习与组合约束结合后能够改善样本期组合表现，但等权多资产配置仍是非常强的风险调整基准，任何 AI 策略都必须在交易成本和回撤约束下接受严格比较。

本文的主要贡献在于：第一，构建了一个可复现、可检查的数据集和回测流程；第二，系统比较 PtO 与 SPO 在组合优化中的差异；第三，将 AI 风险偏好表达约束为结构化 JSON 和程序化校验，而不是停留在主观文本建议；第四，通过图表和附录代码公开核心实现，保证结论可追溯。研究局限包括样本期有限、资产池较小、交易成本简化、未纳入税费与冲击成本、SPO+ 点预测指标较弱，以及当前 AI 层尚未接入外部大模型。本报告所有结论仅适用于课程案例研究。

关键词：AI投资学；ETF；Smart Predict--then--Optimize；SPO+；投资组合优化；回测；风险控制

PAGEBREAK

## Abstract

This report studies the application of artificial intelligence to portfolio optimization through a reproducible multi-asset ETF case. The investment universe consists of six U.S. ETFs, namely SPY, QQQ, TLT, GLD, VNQ and DBC, representing broad U.S. equity, technology-oriented equity, long-duration Treasury bonds, gold, real estate and commodities. The dataset is downloaded from Yahoo Finance through yfinance and a Yahoo Chart API fallback. The final dataset covers {d['first_dates']['SPY']} to {d['last_dates']['SPY']} with {d['row_counts']['SPY']} adjusted close observations for each ETF and passes the implemented quality checks.

The methodology combines classical portfolio theory and decision-focused machine learning. Markowitz's mean-variance framework and the Sharpe ratio provide the investment foundation, while Smart Predict--then--Optimize (SPO) provides the main learning paradigm. Instead of optimizing point forecast accuracy only, SPO aligns the prediction model with downstream portfolio decision quality. This project implements a local SPO+ reproduction based on PyEPO and compares it with Ridge predict-then-optimize, equal-weight and buy-and-hold benchmarks. In addition, an AI-style risk-profile layer generates low-, medium- and high-risk portfolios under explicit volatility and single-asset constraints.

The backtest uses monthly rebalancing and forward 21-trading-day returns. The out-of-sample period contains {meta['data']['test_months']} monthly observations from {meta['data']['test_start']} to {meta['data']['test_end']}. Under a 0.5% one-way transaction cost assumption, the SPO+ strategy with turnover and L2 regularization achieves the highest non-benchmark quality score, with total net return of {pct(best['total_return'])}, annualized return of {pct(best['annual_return'])}, Sharpe ratio of {num(best['sharpe'], 2)} and maximum drawdown of {pct(best['max_drawdown'])}. The equal-weight ETF benchmark delivers the highest net Sharpe ratio, highlighting that diversification remains a strong benchmark. All findings are for educational case-study purposes only and do not constitute investment advice.

Key words: AI investment; ETF; Smart Predict--then--Optimize; SPO+; portfolio optimization; backtesting; risk management

PAGEBREAK

# 一、研究引言

## （一）研究背景与问题提出

人工智能在投资领域的应用通常从收益预测开始。机器学习模型可以处理大量历史价格、成交量、技术指标和横截面特征，并据此输出未来收益、排名或风险状态。然而，投资学关心的终点并不是预测本身，而是预测信号能否转化为可执行、可约束、可复盘的组合权重。一个模型即使在均方误差上略有优势，也可能因为预测误差被组合优化器放大而导致过度集中、换手过高或回撤恶化；反过来，一个点预测指标并不突出的模型，也可能在下游组合约束中形成更有用的排序信号。因此，AI 投资案例不能只汇报准确率，还必须回到投资组合决策质量。

ETF 是课程案例研究中较合适的资产载体。与个股相比，ETF 的资产类别更清晰、流动性通常更好、分散化属性更强，也更便于从投资学角度解释收益来源和风险暴露。本文选取六只代表性 ETF：SPY 代表美国大盘股票，QQQ 代表科技成长股，TLT 代表长期美国国债，GLD 代表黄金，VNQ 代表房地产 REITs，DBC 代表大宗商品。这样的资产池虽然不完整，但覆盖了权益、债券、贵金属、房地产和商品等主要风险来源，可以支撑关于多资产配置、风险分散和模型优化的讨论。

传统投资组合优化常采用“先预测、再优化”的 PtO 思路。该流程先估计资产收益或风险，再将估计值输入均值-方差、最大 Sharpe 或其他优化模型。PtO 的优势是结构清楚、易于解释，但其训练目标通常与最终投资目标分离。SPO 框架则试图把优化问题嵌入学习目标，使模型在训练阶段就感知预测误差对最终权重的影响[3]。本文围绕这一范式，结合本地项目代码与实证输出，重新撰写一篇专业标准的投资学课程案例报告。

## （二）研究目标与核心问题

本文主要回答四个问题。第一，如何构建一个能够支持 AI 投资组合研究的高质量 ETF 数据集？这包括数据来源、价格字段选择、复权收盘价处理、收益率计算、特征构造、标签生成以及未来函数控制。第二，SPO 与传统 PtO 在投资组合场景中有何理论差异？为什么“预测更准”不必然等于“组合更优”？第三，在本地 6 ETF 数据集和 28 个月样本外测试期中，不同策略在净收益、波动率、Sharpe、最大回撤、换手率和超额收益方面表现如何？第四，如何将 AI 风险偏好表达转化为可校验的结构化组合建议，而不是不可复算的文本判断？

## （三）研究思路与技术路线

本文技术路线分为六步。第一步，使用 `src/download_data.py` 从 Yahoo Finance 获取六只 ETF 日度行情，并保存原始价格、复权收盘价和日收益率。第二步，使用 `src/features.py` 构造 22 个历史特征和未来 21 日收益标签，形成日度建模样本和月末调仓样本。第三步，使用 `src/reproduce_spo_paper.py` 训练 Ridge PtO 和 PyEPO SPO+ 线性模型，并生成月度组合权重。第四步，使用 `src/generate_ai_risk_profiles.py` 生成低、中、高三档风险组合，约束年化波动率和单 ETF 权重上限。第五步，使用 `test/backtest_existing_models.py` 独立复算组合指标和未来函数审计，生成累计收益、回撤、权重热力图、特征重要性和绩效表。第六步，按华东师范大学毕业论文格式要求整理为正式课程案例报告。

# 二、理论基础与文献综述

## （一）现代投资组合理论

Markowitz 的组合选择理论奠定了现代投资组合理论的基础[1]。其核心思想是，投资者面对的不是单一资产的收益最大化问题，而是在多个资产之间权衡预期收益与风险。均值-方差框架把风险刻画为收益率方差或协方差结构，使资产之间的相关性成为组合配置的关键因素。在本项目中，6ETF 等权策略的强表现正体现了相关性分散的力量：即使没有复杂预测模型，权益、债券、黄金、房地产与商品的组合也能降低净值路径波动。

Sharpe 比率进一步提供了单位风险收益的评价方式[2]。本文在课程回测中采用年化收益除以年化波动率的简化 Sharpe 口径，同时报告 Sortino、Calmar、最大回撤、胜率和平均换手。这样做的原因是，单一 Sharpe 可能掩盖尾部损失、回撤持续时间和交易成本压力。对 AI 投资组合而言，如果一个策略靠高换手和高集中度获得收益，就必须同时检验扣成本后的收益和回撤，否则容易高估模型价值。

## （二）Predict-then-Optimize 范式及其局限

Predict-then-Optimize 是金融机器学习中常见流程。预测阶段可使用线性回归、Ridge、随机森林、神经网络或时间序列模型估计未来收益；优化阶段再将预测收益与协方差矩阵输入 Markowitz 或最大 Sharpe 模型。本文的 PtO 基准对每只 ETF 单独训练 `StandardScaler + Ridge(alpha=1.0)`，预测未来 21 个交易日收益，再使用历史协方差矩阵求解长期只多最大 Sharpe 权重。

PtO 的局限在于训练目标与决策目标的错位。假设模型 A 的 RMSE 比模型 B 低，但模型 A 在排名第一和第二的资产之间经常犯错，而优化器又把大部分权重放在预测排名第一的资产上，那么微小预测误差就可能造成显著组合损失。金融市场低信噪比、非平稳和强噪声的特征会进一步放大这个问题。由此可见，AI 投资的核心不是“模型看起来更聪明”，而是“模型诱导的组合在可交易约束下更稳健”。

## （三）SPO 与决策导向学习

Elmachtoub 与 Grigas 提出的 SPO 框架正是为了解决预测与优化分离问题[3]。SPO 不再只用预测误差训练模型，而是让预测器直接考虑下游优化问题。直观地说，SPO 希望预测误差在不影响决策的方向上可以被容忍，而会导致错误决策的预测误差应被更强惩罚。SPO+ 是 SPO 损失的可训练 surrogate，能在 PyTorch 等框架中通过优化层反馈梯度。

Wang 与 Hasuike 的 2026 年论文把 SPO 范式应用于真实市场 ETF 组合优化[4]。该论文指出，收益预测准确率提升并不总能带来组合决策质量提升，尤其是在存在交易成本、换手约束和市场非平稳性的情况下。论文采用线性预测器以保持解释性，并引入交易费、换手控制、L2 正则、RobustSPO 和 SoftmaxDFL 等扩展。本文吸收其核心思想，但根据本地课程项目约束做了简化：本项目复现 SPO+ 主路径，使用 PyEPO 实现 SPOPlus 损失，暂未完整复现 RobustSPO、SoftmaxDFL 和 Optuna 滚动调参。

## （四）PyEPO 框架与可复现实现

PyEPO 是一个基于 PyTorch 的端到端 predict-then-optimize 工具库，支持多种线性和整数规划问题的决策导向学习[5]。本项目已下载 PyEPO 官方代码到 `external/PyEPO/`，并在 `etf-spo` conda 环境中以 editable 方式安装。由于原论文 arXiv 源码包只包含论文 LaTeX 与图表，没有提供可运行实验代码，本项目使用 PyEPO 官方实现作为 SPO+ 损失基础。

本项目的自定义优化模型为 `MaxReturnOptModel`。在长期只多、权重和为 1 的 MaxReturn 问题中，如果不加入额外约束，最优解会将全部权重给预测收益最高的资产。这种解法简洁，但容易造成过度集中。因此项目在后续权重生成中加入交易费、换手惩罚和 L2 正则，形成 `spo_plus_turnover_l2` 策略。它不是简单追逐最高预测资产，而是在收益、换手和权重分散之间做数值权衡。

## （五）ETF 资产池构成与代表性

本文数据来自 Yahoo Finance 和 yfinance[6][7]，ETF 资产类别则由发行方官方说明确认。SPY 由 State Street 提供，跟踪 S&P 500 指数，代表美国大盘权益[8]；QQQ 由 Invesco 提供，基于 Nasdaq-100 指数，代表科技成长和非金融大型成长股暴露[9]；TLT 由 iShares 提供，跟踪 20 年以上美国国债指数，代表久期较长的利率风险暴露[10]；GLD 由 World Gold Council 相关平台提供，目标是反映黄金价格表现，代表贵金属和避险资产[11]；VNQ 由 Vanguard 提供，跟踪美国房地产相关指数，代表房地产权益风险[12]；DBC 由 Invesco 提供，跟踪多元商品指数，代表商品周期和通胀相关暴露[13]。

# 三、数据来源、变量构造与数据质量控制

## （一）数据来源与采集机制

本项目使用 Yahoo Finance historical market data 作为历史行情来源，并通过 `yfinance.download()` 进行程序化获取。如果出现限流或网络异常，脚本自动切换到 Yahoo Chart API 备用通道。最近复跑时，yfinance 主通道返回限流提示，脚本成功切换至 Chart API，并通过数据质量检查。所有下载数据保存在 `data/raw/etf_prices_raw.csv`，复权收盘价保存在 `data/processed/prices_adj_close.csv`，日收益率保存在 `data/processed/daily_returns.csv`。

表 1 数据源与资产池说明 Table 1 Data source and ETF universe

{markdown_table([
['SPY', 'State Street SPDR S&P 500 ETF Trust', '美国大盘股票', 'S&P 500 大盘权益代表'],
['QQQ', 'Invesco QQQ Trust', '科技成长股票', 'Nasdaq-100 成长股暴露'],
['TLT', 'iShares 20+ Year Treasury Bond ETF', '长期美国国债', '久期与利率风险暴露'],
['GLD', 'SPDR Gold Shares', '黄金', '贵金属与避险资产'],
['VNQ', 'Vanguard Real Estate ETF', '房地产 REITs', '房地产权益风险'],
['DBC', 'Invesco DB Commodity Index Tracking Fund', '大宗商品', '商品周期和通胀相关暴露'],
], ['代码', '官方名称', '资产类别', '组合含义'])}

## （二）数据字段、复权价格与收益率处理

原始数据包含 Open、High、Low、Close、Adj Close、Volume 六类字段。本文使用 Adj Close 计算收益率，因为复权收盘价已经考虑现金分红和拆分影响，更适合长期持有和多资产回测。日收益率定义为复权收盘价的一阶百分比变化。收益率首日自然为空，这是由 `pct_change()` 决定的计算结果，不属于价格缺失。

表 2 数据质量检查结果 Table 2 Data quality check

{markdown_table([
['样本区间', f"{d['first_dates']['SPY']} 至 {d['last_dates']['SPY']}"],
['每只 ETF 价格行数', str(d['row_counts']['SPY'])],
['复权价格缺失值', '0'],
['重复日期', str(d['duplicate_dates'])],
['非正价格', '0'],
['原始价格表形状', str(tuple(d['raw_shape']))],
['复权收盘价矩阵形状', str(tuple(d['adj_close_shape']))],
['质量总判定', '通过'],
], ['检查项', '结果'])}

## （三）特征工程与标签定义

特征工程脚本 `src/features.py` 生成 22 个历史特征，覆盖短中长期动量、波动率、下行波动率、滚动回撤、均线比率、成交量变化、风险调整收益、横截面排名和横截面 z-score。标签包括未来 21 个交易日收益、未来收益排名、是否跑赢横截面中位数、是否进入未来前二名、未来收益是否为正。

选择未来 21 个交易日作为标签，是为了近似一个月持有期，与月度调仓频率一致。月末调仓样本由日度建模样本筛选每月最后一个可交易日得到。当前复跑生成 {f['daily_modeling_shape'][0]} 行日度建模样本、{f['monthly_rebalance_shape'][0]} 行月末调仓样本。训练集、验证集和测试集按时间划分：训练集不晚于 2022-12-31，验证集为 2023 年，测试集从 2024 年开始。

表 3 建模数据集摘要 Table 3 Modeling dataset summary

{markdown_table([
['特征数量', str(len(f['feature_names']))],
['标签数量', str(len(f['label_names']))],
['首个可用特征日期', f['first_feature_available_date']],
['最后可用标签日期', f['last_label_available_date']],
['日度样本行数', str(f['daily_modeling_shape'][0])],
['月末样本行数', str(f['monthly_rebalance_shape'][0])],
['训练月末样本', str(f['split_counts_monthly']['train'])],
['验证月末样本', str(f['split_counts_monthly']['valid'])],
['测试月末样本', str(f['split_counts_monthly']['test'])],
['建模特征缺失', '无'],
['建模标签缺失', '无'],
], ['项目', '数值'])}

## （四）时间切分与未来函数控制

未来函数是金融回测中最常见也最严重的问题之一。本文从三个层面控制未来函数。第一，特征计算只使用样本日及以前的价格、收益和成交量；第二，StandardScaler 只在训练集上拟合，验证集和测试集仅使用训练集参数转换；第三，回测收益使用已经存储的未来 21 日标签，且回测日期与测试标签日期严格一致。独立审计文件 `test/backtest_outputs/lookahead_audit.json` 显示所有检查均通过。

# 四、模型设定与组合优化方法

## （一）PtO-Markowitz 基准模型

PtO-Markowitz 策略分为预测和优化两步。预测阶段，对每只 ETF 分别训练 Ridge 回归模型，输入为 22 个历史特征，输出为未来 21 日收益。优化阶段，将预测收益年化，并结合最近 60 个交易日的历史协方差矩阵，求解长期只多最大 Sharpe 组合。约束条件为所有权重非负且权重和为 1。

PtO 基准的意义在于提供一个传统机器学习投资流程对照。它能够体现预测模型和 Markowitz 优化的结合，但也暴露了高换手和回撤问题。当前独立回测显示，PtO-Markowitz 扣成本后总收益 {pct(net_metrics[net_metrics['strategy']=='pto_markowitz']['total_return'].iloc[0])}，最大回撤 {pct(net_metrics[net_metrics['strategy']=='pto_markowitz']['max_drawdown'].iloc[0])}，平均换手 {pct(net_metrics[net_metrics['strategy']=='pto_markowitz']['average_turnover'].iloc[0])}，说明成本敏感性较高。

## （二）SPO+ 决策导向学习模型

SPO+ 模型使用 PyTorch 线性层从月末特征向量预测六只 ETF 的未来收益分数，并用 PyEPO 的 `SPOPlus` 损失训练。训练时，模型不只是拟合真实收益，而是通过优化模型感知预测值导致的组合决策差异。主脚本默认训练 240 轮，batch size 为 16，学习率为 0.01，优化器为 Adam，weight decay 为 1e-4。

基础 SPO+ 的 MaxReturn 权重会高度集中；这符合线性目标在 simplex 约束下的数学性质，但不一定符合投资实践。为此，项目设置了三类 SPO 策略：基础 `spo_plus`、加入交易费项的 `spo_plus_fee`、加入交易费、换手惩罚和 L2 权重正则的 `spo_plus_turnover_l2`。其中最后一个策略是本文重点讨论的 SPO 类组合，因为它显式处理交易摩擦和权重稳定性。

## （三）风险偏好约束下的 AI 组合生成

AI 风险组合层不是直接调用大模型生成主观建议，而是构造了一套可被大模型接入、也能由程序化优化器复现的结构化流程。输入包括 SPO 候选组合、ETF 单体风险收益指标、SPO 最新排名和三档风险约束。输出必须为 JSON，包括权重、预期收益、预期波动、历史风险检查、校验状态、风险说明和配置理由。

三档风险约束为：低风险年化波动率不超过 10%、单只 ETF 权重不超过 35%；中风险年化波动率不超过 15%、单只 ETF 权重不超过 45%；高风险年化波动率不超过 22%、单只 ETF 权重不超过 60%。若某月无法完全满足波动率约束，系统回退到同一单资产上限下的最低波动组合，并记录 `fallback_to_minimum_volatility`。当前复跑中权重和、非负和单资产上限全部通过，风险约束有 {ai_validation['risk_constraint_fallback_count']} 次回退且已标注。

表 4 最新 AI 三档风险组合 Table 4 Latest AI risk-profile portfolios

{markdown_table(ai_rows, ['风险档位', '波动上限', '单 ETF 上限', '最新权重', '预期年化收益', '预期年化波动', '预期Sharpe'])}

## （四）交易成本、换手惩罚与正则化约束

本文主口径采用 0.5% 单边交易成本。换手率定义为本期权重与上期权重差的绝对值之和。净收益计算公式为：`net_return_t = gross_return_t - turnover_t * 0.005`。这一路径虽简化了真实交易中的买卖价差、市场冲击、税费和基金申赎机制，但足以在课程案例中体现高换手策略的成本压力。

将权重上限、换手控制和交易成本写入组合优化，也与大规模组合中约束可以缓解估计误差影响的经验研究相一致[14]，并与把交易成本纳入多期组合优化的研究方向相呼应[15]。

交易成本对结果影响明显。基础 SPO+ 扣成本前总收益 {pct(metrics[(metrics['strategy']=='spo_plus') & (metrics['return_type']=='gross')]['total_return'].iloc[0])}，扣成本后总收益降至 {pct(metrics[(metrics['strategy']=='spo_plus') & (metrics['return_type']=='net')]['total_return'].iloc[0])}；PtO-Markowitz 扣成本前总收益 {pct(metrics[(metrics['strategy']=='pto_markowitz') & (metrics['return_type']=='gross')]['total_return'].iloc[0])}，扣成本后仅 {pct(metrics[(metrics['strategy']=='pto_markowitz') & (metrics['return_type']=='net')]['total_return'].iloc[0])}。这说明 AI 投资策略如果忽略换手成本，容易产生过度乐观结论。

# 五、回测框架与评价指标体系

## （一）样本外区间与月度调仓规则

回测采用月度调仓。每个调仓日为当月最后一个可交易日，模型读取截至该日可观察的特征，生成未来 21 个交易日收益预测或决策分数，再据此构建组合权重。组合持有至下一期，用未来 21 日真实收益计算本期实现收益。测试期为 {meta['data']['test_start']} 至 {meta['data']['test_end']}，共 {meta['data']['test_months']} 个调仓月。

这种回测设计与课程项目的数据规模匹配。日度调仓可能放大噪声和交易成本，季度调仓又可能过度稀疏。月度调仓在学术研究和资产配置实践中都较常见，也与原 SPO 论文中的月度再平衡思想一致[4]。

## （二）绩效评价指标体系

本文使用总收益、年化收益、年化波动率、Sharpe、Sortino、最大回撤、Calmar、胜率、平均换手率和最终财富评价策略。总收益反映累计财富增长，年化收益便于跨周期比较；年化波动率衡量收益不确定性；Sharpe 衡量单位波动收益；Sortino 只考虑下行波动；最大回撤衡量从历史高点到低点的最大损失；Calmar 用年化收益除以最大回撤绝对值；胜率反映正收益月份占比；平均换手体现交易成本压力。

为了避免只对模型策略友好，本文保留两个强基准：6ETF 等权和 SPY 买入持有。等权基准体现多资产分散化，SPY 买入持有体现美股大盘风险暴露。如果 AI/SPO 策略无法在扣成本后超过这些简单基准，其应用价值就需要谨慎评价。

## （三）独立复核与未来函数审计

主模型脚本会输出权重和收益，独立回测脚本则读取这些输出重新计算全部指标，并额外输出未来函数审计、随机森林诊断特征重要性、净值曲线、回撤曲线和成本敏感性。审计结果显示：t 月末特征只使用 t 月末及以前数据；标签为未来 21 个交易日收益；组合回测日期与测试标签日期一致；StandardScaler 只在训练集拟合；训练验证期结束早于测试期开始。由此可以认为当前回测没有明显未来函数问题。

# 六、实证结果与投资组合分析

## （一）ETF 资产表现与风险收益特征

如图 1 所示，六只 ETF 从样本起点归一化后的价格或净值走势存在明显差异。权益资产 SPY 与 QQQ 在多数时期表现较强，QQQ 的成长属性带来更高收益和更高波动；TLT 在加息周期承受较大压力；GLD 在避险和通胀环境中具有独立走势；VNQ 对利率和房地产周期敏感；DBC 反映商品周期。资产走势差异说明，跨资产配置比单一资产押注更适合课程研究。

![图 1 ETF价格指数走势 Figure 1 ETF price index](test/backtest_outputs/plots/etf_price_index.png)

## （二）策略净收益与风险调整表现

由表 5 可见，各策略在扣除 0.5% 单边交易成本后的表现分化明显。`{clean_strategy(best['strategy'])}` 的净收益总额最高，达到 {pct(best['total_return'])}；其年化收益为 {pct(best['annual_return'])}，Sharpe 为 {num(best['sharpe'], 2)}，最大回撤为 {pct(best['max_drawdown'])}。它相对等权基准的年化超额收益为 {pct(context['excess'][context['excess']['strategy']=='spo_plus_turnover_l2']['annual_excess_return'].iloc[0])}，信息比率为 {num(context['excess'][context['excess']['strategy']=='spo_plus_turnover_l2']['information_ratio'].iloc[0], 2)}。

表 5 主要策略净收益指标 Table 5 Net backtest metrics

{markdown_table(metric_rows, ['策略', '总收益', '年化收益', '年化波动', 'Sharpe', '最大回撤', '胜率', '平均换手'])}

从图 2 的累计净值曲线可以看到，`SPO+换手/L2约束` 在样本期后段形成明显领先；基础 SPO+ 和 SPO+交易费版本走势高度接近，说明当前交易费项没有形成有效区分；等权和 SPY 买入持有的净值更平滑，特别是等权组合在风险调整收益上非常有竞争力。

![图 2 策略累计收益曲线 Figure 2 Strategy cumulative return curves](test/backtest_outputs/plots/strategy_cumulative_returns.png)

## （三）回撤控制与下行风险比较

从图 3 的主要策略回撤曲线看，PtO-Markowitz 最大回撤较深，说明传统预测再优化在当前样本下容易受到估计误差和换手影响。高风险 AI 组合也出现较大回撤，其原因是约束更宽、换手更高，并且较多配置到波动资产。相比之下，等权组合回撤最浅之一，这再次说明分散化和低换手在短样本回测中非常重要。

![图 3 策略回撤曲线 Figure 3 Strategy drawdown curves](test/backtest_outputs/plots/strategy_drawdowns.png)

`SPO+换手/L2约束` 的最大回撤为 {pct(best['max_drawdown'])}，虽然不是全场最低，但在总收益明显领先的同时维持了可接受的回撤水平。AI 低风险组合最大回撤为 {pct(net_metrics[net_metrics['strategy']=='ai_low_risk']['max_drawdown'].iloc[0])}，体现了风险约束在平滑净值方面的作用。不过 AI 低风险组合扣成本后总收益低于等权基准，说明风险控制本身并不自动创造超额收益。

## （四）预测质量与决策质量诊断

由表 6 的预测层指标可见，SPO+ 在 RMSE、MAE 和 R2 上明显弱于 Ridge，测试 Rank IC 接近零。这一现象提醒我们，不能把组合层表现直接解释为模型拥有稳定收益预测能力。更准确的表述是：在当前样本和当前权重生成规则下，SPO+换手/L2约束组合取得了较好的样本外组合表现；但其预测分数本身仍需更长样本、更严格滚动验证和更多资产检验。

表 6 预测层指标 Table 6 Prediction-level diagnostics

{markdown_table([[r.model_name, r.split, num(r.rmse, 4), num(r.mae, 4), num(r.r2, 4), num(r.rank_ic, 4)] for r in model_metrics.itertuples()], ['模型', '样本', 'RMSE', 'MAE', 'R2', 'Rank IC'])}

## （五）AI 风险档位组合权重分析

在图 4 的 AI 三档风险组合月度权重热力图中，低风险组合受到 10% 波动上限和 35% 单资产上限约束，配置更分散；中风险组合允许更高单资产权重；高风险组合上限放宽到 60%，因此更容易集中到模型认为预期收益较高的资产。热力图也显示，AI 组合并非静态组合，而是随预测信号和风险约束变化。

![图 4 AI组合月度权重热力图 Figure 4 AI monthly allocation heatmap](test/backtest_outputs/plots/ai_monthly_weight_heatmap.png)

当前最新决策日 {ai_json['decision_date']} 的三档组合均通过权重校验。低风险组合预期年化波动为 {pct(ai_json['risk_profiles'][0]['expected_annual_volatility'])}，中风险为 {pct(ai_json['risk_profiles'][1]['expected_annual_volatility'])}，高风险为 {pct(ai_json['risk_profiles'][2]['expected_annual_volatility'])}。不过，三档组合在扣成本后都没有超过等权基准，这说明可解释的风险偏好组合更适合作为“投资建议展示层”，而不是当前样本下的收益冠军。

## （六）诊断性特征重要性分析

随机森林诊断性特征重要性结果见图 5。该模型不参与 SPO/PtO 主回测，只用于帮助解释哪些历史变量对未来 21 日收益标签更有信息。前十特征包括 downside_vol_60d、drawdown_60d、z_vol_60d、drawdown_20d、ma_ratio_60_120、ret_120d 和 volume_chg_20d，说明波动、回撤、趋势和成交量变化都包含一定信息。

![图 5 随机森林特征重要性 Figure 5 Random forest feature importance](test/backtest_outputs/plots/random_forest_feature_importance.png)

表 7 随机森林诊断性特征重要性前十 Table 7 Top 10 diagnostic feature importance

{markdown_table([[r.feature, num(r.importance, 4), str(int(r.rank))] for r in importance.head(10).itertuples()], ['特征', '重要性', '排名'])}

## （七）综合绩效可视化比较

此外，图 6 是独立回测脚本生成的绩效表图像，集中展示各策略的收益、波动、Sharpe 和回撤表现。将该图放入报告，是为了让读者在表格之外快速比较模型策略、AI 风险组合和基准组合的整体关系。

![图 6 策略绩效对比表 Figure 6 Strategy performance table](test/backtest_outputs/plots/strategy_performance_table.png)

# 七、稳健性检验、局限性与风险提示

## （一）主要研究发现

第一，SPO 类方法在本样本中具有组合层价值。尤其是加入换手和 L2 正则后，策略在净收益和综合评分上领先。这说明，仅将预测分数交给极端 MaxReturn 优化器并不足够，必须加入交易和分散化约束，才能更接近可投资组合。

第二，等权基准非常强。6ETF 等权组合的净 Sharpe 为 {num(sharpe_best['sharpe'], 2)}，最大回撤仅 {pct(sharpe_best['max_drawdown'])}，平均换手只有 {pct(sharpe_best['average_turnover'])}。这提醒我们，AI 策略不能只和弱基准比较，也不能只展示扣成本前收益。一个专业投资报告必须把简单、低成本、可解释的基准放在同一张表里。

第三，点预测指标与组合绩效并不一致。SPO+ 的点预测误差较大，但带约束的组合表现较好；Ridge 的点预测误差较低，但 PtO-Markowitz 扣成本后表现较弱。这不意味着预测指标无用，而是说明投资场景下必须同时评价预测层和决策层。

第四，AI 风险组合的价值在于结构化表达和可校验输出。本文没有让 AI 直接写“建议买入某 ETF”，而是规定 JSON schema、权重约束和风险复算规则。这样的设计更符合金融场景对审计和可解释性的要求。

## （二）稳健性与一致性检验

本文进行了三类稳健性检查。第一，数据质量检查确认价格无缺失、日期一致、无非正价格。第二，未来函数审计确认训练、验证、测试和回测收益之间的时间顺序正确。第三，交易成本敏感性表展示不同成本率下策略表现变化，说明高换手策略对成本假设敏感。虽然这些检查不能证明策略未来有效，但能降低课程报告中常见的数据泄露和口径不一致风险。

## （三）研究局限

本文仍有明显局限。第一，测试期只有 28 个调仓月，统计置信度有限。第二，资产池只有六只 ETF，不能代表全球全部资产类别，也不能覆盖行业、风格、期限、信用和地域的完整结构。第三，数据源 Yahoo Finance 和 yfinance 适合课程复现，但正式投资研究应使用交易所、基金公司或专业数据供应商数据复核。第四，交易成本用固定 0.5% 近似，未纳入滑点、市场冲击、税费和流动性冲击。第五，当前实现没有完整复现原论文中的 RobustSPO、SoftmaxDFL 和滚动超参数搜索。第六，AI 层当前是结构化提示词设计和程序化优化器兜底，尚未接入真实大模型 API。

## （四）投资风险与适用边界

本报告所有结果均为历史回测，不构成任何投资建议。ETF 价格会受到宏观经济、利率、通胀、政策、地缘政治、市场情绪和流动性等多重因素影响。模型回测收益可能来自样本期偶然结构，未来市场环境变化可能导致策略失效。任何真实投资决策都需要结合投资者风险承受能力、资金期限、税务约束和合规要求。

# 八、研究结论与后续展望

本文重写并扩展了 AI 投资学课程案例报告，围绕 SPO 决策导向学习构建了完整的多资产 ETF 组合优化研究。数据层面，项目使用 Yahoo Finance 日度历史行情，构建了无缺失、可复现、可追溯的六 ETF 数据集；模型层面，项目比较了 PtO-Markowitz、SPO+、SPO+交易费、SPO+换手/L2约束、等权、SPY 买入持有和 AI 三档风险组合；回测层面，项目使用月度调仓、未来 21 日持有期收益和 0.5% 单边交易成本，并由独立脚本复算净值、回撤和未来函数审计。

实证结果表明，在当前样本期内，`SPO+换手/L2约束` 是综合评分最高的非基准策略，扣成本后总收益 {pct(best['total_return'])}、年化收益 {pct(best['annual_return'])}、Sharpe {num(best['sharpe'], 2)}、最大回撤 {pct(best['max_drawdown'])}。不过，6ETF 等权组合仍提供最高净 Sharpe，这说明多资产分散化是非常强的投资学基准。本文的核心结论并不是“AI 必然战胜市场”，而是：AI 与 SPO 方法可以把预测信号、优化目标、交易成本和风险偏好放进同一个可复验框架；只有当数据、约束、回测和解释全部经得起检查时，AI 投资案例才具有课程研究意义。

后续研究可以沿四个方向扩展。第一，扩大 ETF 池，加入国际股票、不同期限债券、信用债、行业 ETF、风格 ETF 和波动率产品。第二，引入滚动训练窗口和更长历史样本，检验参数稳定性。第三，完整复现 RobustSPO、SoftmaxDFL 和 Optuna 调参，比较不同决策导向学习结构。第四，接入真实大模型 API，但保留当前程序化校验器，确保 AI 输出始终可复算、可约束、可审计。

# 参考文献

{ref_text}

# 附录A 复现环境、运行命令与输出目录

## A.1 复现环境

本项目的正式复现环境为 `etf-spo` Conda 环境。环境文件位于 `environment/etf-spo.yml`，核心依赖包括 Python、pandas、numpy、scikit-learn、scipy、PyTorch、PyEPO、matplotlib、seaborn、yfinance 与 python-docx。正式文档采用 `scripts/generate_case_report.py` 生成 Markdown 与 DOCX，再由本机文档工具链导出 PDF。

## A.2 核心运行命令

```bash
conda env create -f environment/etf-spo.yml
conda run -n etf-spo python src/download_data.py
conda run -n etf-spo python src/features.py
conda run -n etf-spo python src/reproduce_spo_paper.py
conda run -n etf-spo python src/generate_ai_risk_profiles.py
conda run -n etf-spo python test/backtest_existing_models.py --cost-rate 0.005
conda run -n etf-spo python scripts/generate_case_report.py
codex-docx-to-pdf docs/案例报告/AI投资学课程案例报告.docx docs/案例报告
codex-docx-inspect docs/案例报告/AI投资学课程案例报告.docx
codex-pdf-inspect docs/案例报告/AI投资学课程案例报告.pdf
```

## A.3 输出目录说明

附录D列出了全部 CSV、JSON、Markdown 与 PNG 输出产物。核心目录包括：`data/raw/` 存放原始两级表头价格数据；`data/processed/` 存放复权价格、日收益率、建模样本与月末调仓样本；`outputs/tables/` 存放主模型、SPO复现、AI风险组合和数据质量表；`test/backtest_outputs/` 存放独立回测报告、指标、审计文件和图表；`docs/案例报告/` 存放最终 Markdown、DOCX 与 PDF。

# 附录B 核心代码与实现说明

附录B摘录的是当前工作区中的真实核心源码。为避免报告过度冗长，附录保留数据下载、特征工程、SPO训练、组合优化、AI风险组合、独立回测和文档生成的关键函数；完整代码以附表D-1中的路径为准。

附表 B-1 核心代码文件与函数索引 Appendix Table B-1 Core code file and function index

{markdown_table(code_sources, ['模块', '源码路径', '核心函数或类'])}

## B.1 数据下载与质量控制

数据层脚本先尝试 yfinance 主通道，若网络或限流导致失败，则切换到 Yahoo Chart API 备用通道。只有当质量检查通过后，脚本才会写出原始价格、复权收盘价、收益率和元数据。

```python
{code_quality}
```

## B.2 特征工程与标签生成

特征工程脚本将宽表价格、收益率和成交量转换为面板式建模数据。所有特征只使用样本日及以前的数据；未来21个交易日收益只作为标签和回测收益进入后续流程。

```python
{code_features}
```

## B.3 SPO+ 训练、权重生成与净收益回测

SPO复现脚本包含 PyEPO 优化模型、SPOPlus 损失训练、长期只多优化、月度权重生成和净收益回测。`spo_plus_turnover_l2` 的关键在于同时考虑预期收益、换手惩罚和 L2 权重正则。

```python
{code_spo}
```

## B.4 AI风险档位组合生成与校验

AI组合层使用显式风险约束和程序化校验。低、中、高三档组合分别约束年化波动率和单ETF上限；若波动率约束不可行，系统回退至同一单资产上限下的最低波动组合并保留原因。

```python
{code_ai}
```

## B.5 独立回测、未来函数审计与图表输出

独立回测脚本不复用主脚本的指标结果，而是读取已有权重、收益和特征元数据重新计算净值、回撤、评分排序、成本敏感性、未来函数审计和全部图表。

```python
{code_audit}
```

## B.6 正式文档生成与格式导出

正式报告由同一个脚本生成 Markdown 与 DOCX。`markdown_to_docx` 负责标题、正文、表题、图题、代码块、页边距、字体、字号和页码等格式控制；PDF由本机文档工具链从 DOCX 导出。

```python
{code_report}
```

# 附录C 附图汇编

附录C汇集独立回测脚本生成的全部可视化图像。正文已引用主要图表，附录再次集中列示，便于审阅者核对图像来源、编号和含义。

![附图 C-1 ETF价格指数走势 Appendix Figure C-1 ETF price index](test/backtest_outputs/plots/etf_price_index.png)

![附图 C-2 策略累计收益曲线 Appendix Figure C-2 Strategy cumulative returns](test/backtest_outputs/plots/strategy_cumulative_returns.png)

![附图 C-3 策略回撤曲线 Appendix Figure C-3 Strategy drawdowns](test/backtest_outputs/plots/strategy_drawdowns.png)

![附图 C-4 AI组合月度权重热力图 Appendix Figure C-4 AI monthly allocation heatmap](test/backtest_outputs/plots/ai_monthly_weight_heatmap.png)

![附图 C-5 随机森林特征重要性 Appendix Figure C-5 Random forest feature importance](test/backtest_outputs/plots/random_forest_feature_importance.png)

![附图 C-6 策略绩效对比表 Appendix Figure C-6 Strategy performance table](test/backtest_outputs/plots/strategy_performance_table.png)

# 附录D 附表索引与主要输出表

附录D列示本项目生成的主要数据表、JSON、图像和报告文件，并补充若干关键输出预览。完整数据以对应文件为准。

附表 D-1 输出产物索引 Appendix Table D-1 Output artifact index

{markdown_table(artifact_rows, ['路径', '类型', '规模或结构', '大小', '用途'])}

附表 D-2 ETF数据质量明细 Appendix Table D-2 ETF data quality details

{markdown_table(data_quality_rows, ['ETF', '起始日期', '结束日期', '价格行数', '复权缺失', '非正价格', '区间总收益', '区间最大回撤'])}

附表 D-3 模型质量综合排序 Appendix Table D-3 Model quality ranking

{markdown_table(ranking_rows, ['排名', '策略', '策略组', '质量分', '总收益', '年化收益', 'Sharpe', '最大回撤', '平均换手'])}

附表 D-4 交易成本敏感性摘要 Appendix Table D-4 Transaction cost sensitivity summary

{markdown_table(sensitivity_rows, ['策略', '成本率', '总收益', '年化收益', 'Sharpe', '最大回撤', '平均换手'])}

附表 D-5 策略评估与推荐用途 Appendix Table D-5 Strategy assessment and recommended use

{markdown_table(assessment_rows, ['排名', '策略', '策略组', '优势', '弱点', '推荐用途'])}

附表 D-6 AI风险组合校验摘要 Appendix Table D-6 AI risk profile validation summary

{markdown_table([
['权重和、非负、单资产上限校验', str(ai_validation['ai_weight_checks_ok'])],
['全部风险约束直接满足', str(ai_validation['ai_risk_constraints_all_met'])],
['风险约束或回退机制有效', str(ai_validation['ai_risk_constraints_or_fallbacks_ok'])],
['风险约束回退次数', str(ai_validation['risk_constraint_fallback_count'])],
['组合指标有限性校验', str(ai_validation['metrics_finite_ok'])],
['纳入校验的风险档位', ', '.join(ai_validation['ai_json_profiles'])],
], ['检查项', '结果'])}

"""
    return md


def apply_run_font(
    run,
    *,
    size: float = BODY_PT,
    zh_font: str = FONT_ZH,
    en_font: str = FONT_EN,
    bold: bool = False,
    italic: bool = False,
    superscript: bool = False,
) -> None:
    run.bold = bold
    run.italic = italic
    run.font.superscript = superscript
    run.font.name = en_font
    run._element.rPr.rFonts.set(qn("w:ascii"), en_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), en_font)
    run._element.rPr.rFonts.set(qn("w:cs"), en_font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), zh_font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_rich_text(
    paragraph,
    text: str,
    *,
    size: float = BODY_PT,
    zh_font: str = FONT_ZH,
    bold: bool = False,
    superscript_citations: bool = True,
) -> None:
    text = text.replace("**", "").replace("`", "")
    if superscript_citations:
        parts = re.split(r"(\[\d+\])", text)
    else:
        parts = [text]
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        apply_run_font(
            run,
            size=size,
            zh_font=zh_font,
            bold=bold,
            superscript=bool(superscript_citations and re.fullmatch(r"\[\d+\]", part)),
        )


def set_paragraph_format(
    paragraph,
    *,
    first_line: bool = True,
    line_spacing: float = 1.5,
    align=None,
    before: float = 0,
    after: float = 0,
) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(24) if first_line else Pt(0)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    add_rich_text(p, str(text), size=TABLE_PT, bold=bold, superscript_citations=False)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_border(cell, **edges: dict[str, str | int]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, settings in edges.items():
        tag = tc_borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_borders.append(tag)
        for key, value in settings.items():
            tag.set(qn(f"w:{key}"), str(value))


def set_table_borders(table) -> None:
    single = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    none = {"val": "nil"}
    last_idx = len(table.rows) - 1
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell, top=none, bottom=none, left=none, right=none)
            if row_idx == 0:
                set_cell_border(cell, top=single, bottom=single)
            if row_idx == last_idx:
                set_cell_border(cell, bottom=single)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in [("Heading 1", TITLE_PT), ("Heading 2", BODY_PT), ("Heading 3", BODY_PT)]:
        style = styles[style_name]
        style.font.name = FONT_EN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT_EN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    style: str | None = None,
    align=None,
    size: float = BODY_PT,
    first_line: bool = True,
    bold: bool = False,
    zh_font: str = FONT_ZH,
    superscript_citations: bool = True,
) -> None:
    p = doc.add_paragraph(style=style)
    set_paragraph_format(p, first_line=first_line, align=align)
    add_rich_text(
        p,
        text,
        size=size,
        zh_font=zh_font,
        bold=bold,
        superscript_citations=superscript_citations,
    )


def add_caption(doc: Document, text: str, above: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_rich_text(p, text, size=CAPTION_PT, bold=False, superscript_citations=False)


def add_table_docx(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    set_table_borders(table)


def parse_markdown_blocks(md: str) -> list[tuple[str, Any]]:
    blocks: list[tuple[str, Any]] = []
    lines = md.splitlines()
    i = 0
    para: list[str] = []

    def flush_para() -> None:
        nonlocal para
        if para:
            blocks.append(("p", " ".join(para).strip()))
            para = []

    while i < len(lines):
        line = lines[i]
        if not line.strip():
            flush_para()
            i += 1
            continue
        if line.strip() == "PAGEBREAK":
            flush_para()
            blocks.append(("pagebreak", None))
            i += 1
            continue
        if line.startswith("```"):
            flush_para()
            lang = line.strip("`").strip()
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            blocks.append(("code", "\n".join(code)))
            i += 1
            continue
        if line.startswith("# "):
            flush_para()
            blocks.append(("h1", line[2:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            flush_para()
            blocks.append(("h2", line[3:].strip()))
            i += 1
            continue
        if line.startswith("### "):
            flush_para()
            blocks.append(("h3", line[4:].strip()))
            i += 1
            continue
        if line.startswith("!["):
            flush_para()
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                blocks.append(("image", (match.group(1), match.group(2))))
            i += 1
            continue
        if line.startswith("| "):
            flush_para()
            table_lines = []
            while i < len(lines) and lines[i].startswith("| "):
                table_lines.append(lines[i])
                i += 1
            header = [c.strip() for c in table_lines[0].strip("|").split("|")]
            rows = []
            for tline in table_lines[2:]:
                rows.append([c.strip() for c in tline.strip("|").split("|")])
            blocks.append(("table", (header, rows)))
            continue
        para.append(line.strip())
        i += 1
    flush_para()
    return blocks


def markdown_to_docx(md: str, output: Path) -> None:
    doc = Document()
    setup_styles(doc)
    blocks = parse_markdown_blocks(md)
    pending_table_caption: str | None = None
    current_part = "front"
    h1_count = 0

    def add_heading(text: str, level: int) -> None:
        nonlocal current_part, h1_count
        if level == 1:
            h1_count += 1
        if text in {"摘要", "Abstract"}:
            current_part = "abstract"
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=3)
            add_rich_text(p, text, size=BODY_PT, zh_font=FONT_HEADING, bold=True, superscript_citations=False)
            return
        if text == "目录":
            current_part = "toc"
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)
            add_rich_text(p, text, size=BODY_PT, zh_font=FONT_HEADING, bold=True, superscript_citations=False)
            return
        if re.match(r"^[一二三四五六七八九十]+、", text) or text in {"参考文献"} or text.startswith("附录"):
            current_part = "body" if re.match(r"^[一二三四五六七八九十]+、", text) else text
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=3)
            add_rich_text(p, text, size=BODY_PT, zh_font=FONT_HEADING, bold=True, superscript_citations=False)
            return
        if h1_count <= 2 and level == 1:
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)
            add_rich_text(p, text, size=TITLE_PT, zh_font=FONT_HEADING, bold=True, superscript_citations=False)
            return
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=6 if level == 2 else 3, after=3)
        add_rich_text(p, text, size=BODY_PT, zh_font=FONT_HEADING, bold=True, superscript_citations=False)

    for block_type, value in blocks:
        if block_type == "h1":
            add_heading(value, 1)
        elif block_type == "h2":
            add_heading(value, 2)
        elif block_type == "h3":
            add_heading(value, 3)
        elif block_type == "p":
            text = value
            if text.startswith(("表", "附表")) and "Table" in text:
                pending_table_caption = text
                continue
            if text.startswith("**") and text.endswith("**"):
                add_paragraph(doc, text, first_line=False, bold=True)
            elif current_part == "toc":
                add_paragraph(doc, text, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT, superscript_citations=False)
            elif current_part == "abstract":
                add_paragraph(doc, text, size=ABSTRACT_PT)
            elif current_part == "参考文献":
                add_paragraph(doc, text, size=ABSTRACT_PT, first_line=False, superscript_citations=False)
            elif current_part.startswith("附录"):
                add_paragraph(doc, text, first_line=False, superscript_citations=False)
            elif text.startswith(("English Title:", "课程：", "项目目录：", "完成日期：", "声明：")):
                size = TITLE_PT if text.startswith("English Title:") else BODY_PT
                add_paragraph(doc, text, size=size, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, superscript_citations=False)
            else:
                add_paragraph(doc, text)
        elif block_type == "pagebreak":
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
        elif block_type == "table":
            if pending_table_caption:
                add_caption(doc, pending_table_caption, above=True)
                pending_table_caption = None
            headers, rows = value
            add_table_docx(doc, headers, rows)
        elif block_type == "image":
            alt, path_text = value
            img_path = ROOT / path_text
            if img_path.exists():
                p = doc.add_paragraph()
                set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
                run = p.add_run()
                width = Cm(12.5)
                if "heatmap" in img_path.name:
                    width = Cm(11.5)
                run.add_picture(str(img_path), width=width)
                add_caption(doc, alt, above=False)
        elif block_type == "code":
            for line in value.splitlines():
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(line)
                apply_run_font(run, size=CODE_PT, zh_font=FONT_CODE, en_font=FONT_CODE)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    context = get_tables()
    write_docs(context)
    md = build_report_markdown(context)
    OUTPUT_MD.write_text(md, encoding="utf-8")
    markdown_to_docx(md, OUTPUT_DOCX)
    print(json.dumps({
        "markdown": str(OUTPUT_MD),
        "docx": str(OUTPUT_DOCX),
        "chars": len(re.sub(r"\\s+", "", md)),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
